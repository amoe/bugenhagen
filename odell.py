import docx
import pprint
import random
import lxml
import sys

# python-docx 0.8.7

PATH = "/home/amoe/martin_odell/Object-Oriented Methods a Foundation, Martin and Odell.docx"

doc = docx.Document(PATH)


# public api for document is

 # 'add_heading',
 # 'add_page_break',
 # 'add_paragraph',
 # 'add_picture',
 # 'add_section',
 # 'add_table',
 # 'core_properties',
 # 'element',
 # 'inline_shapes',
 # 'paragraphs',
 # 'part',
 # 'save',
 # 'sections',
 # 'settings',
 # 'styles',
 # 'tables'

# for (index, paragraph) in enumerate(doc.paragraphs):
#     print(index)



all_paragraphs = doc.paragraphs
# index = random.randint(0, len(all_paragraphs))
#index = 3636
index = 256
some_para = all_paragraphs[index]
print(some_para.style.style_id)
# print(some_para.style)

def handle_body_text(para):
    child = lxml.etree.Element('p')
    print(para.paragraph_format.line_spacing)

    

    child.text = para.text
    return [child]

def handle_heading(para):
    child = lxml.etree.Element('h1')
    child.text = para.text
    return [child]

style_handlers = {
    'Style121': handle_heading,
    'Style118': handle_body_text
}

print(len(doc.inline_shapes))

root = lxml.etree.Element('body')

for p in all_paragraphs:
    style_id = p.style.style_id

    handler = style_handlers.get(style_id, lambda x: [])
    elements = handler(p)

    for e in elements:
        root.append(e)

et = lxml.etree.ElementTree(root)
with open('foo.html', 'wb') as f:
    et.write(f, pretty_print=True)

# for paragraph in doc.paragraphs:


# paragraphstyle has ids
# style_id

# i = 0
# for p in all_paragraphs:
#     if p.text == 'A Little History':
#         break

#     i += 1

# print(i)

# API for  paragraph consists of 

 # 'add_run',
 # 'alignment',
 # 'clear',
 # 'insert_paragraph_before',
 # 'paragraph_format',
 # 'part',
 # 'runs',
 # 'style',
 # 'text'

# print(index)
# print(some_para.text)

# for run in some_para.runs:
#     print(run.text)
        # text = ''
        # for run in self.runs:
        #     text += run.text
        # return text



# Run public api:
# You can query for bold, which is the self.font.bold and self.font.italic call
